"""Generate a Word .docx file that lays out all current prompt operators as paper-style boxes.

Usage:
    uv run python scripts/generate_prompt_template_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "文档" / "prompt_template_example.docx"


LEVEL1_COMMON_TASK_GOAL = (
    "You are designing the Level-1 proxy objective function for an Edge-UAV computation "
    "offloading system. For each active task, you must decide whether it should be executed "
    "locally or offloaded to a UAV. The objective must guide Gurobi to make good assignment "
    "decisions while remaining linear in binary variables."
)

LEVEL1_COMMON_ENV = (
    "You are given the current system scale summary, active-task statistics, UAV resource "
    "statistics, the OffloadingModel source code, variable definitions, precomputed constants "
    "(D_hat_local, D_hat_offload, E_hat_comp), and the current iteration summaries task_info "
    "and uav_info."
)

LEVEL1_DECISION_RULES = [
    "Each task must be assigned exactly once within its active window, either locally or to one UAV.",
    "Every task-related summation must filter by self.task[i].active[t]; inactive tasks cannot be used.",
    "The objective must use only supported Gurobi-linear expressions over self.x_local[i, t] and self.x_offload[i, j, t].",
    "Precomputed constants and scalar transforms on constants are allowed, but nonlinear operations on decision variables are not allowed.",
]

LEVEL1_OUTPUT = (
    "Output only a JSON object with exactly two keys: obj_description and obj_code. "
    "The obj_code field must define dynamic_obj_func(self) with valid Python indentation "
    "and a final gb.quicksum(w*c for w, c in zip(costs, weights)) objective."
)

LEVEL2_COMMON_TASK_GOAL = (
    "You are designing the Level-2b trajectory objective weighting strategy for the UAV "
    "trajectory optimizer. The goal is to balance communication-delay surrogate and propulsion "
    "energy while preserving convexity and producing a better downstream evaluation score."
)

LEVEL2_COMMON_ENV = (
    "You are given normalized CVXPY objective components obj_comm_surrogate, obj_propulsion, "
    "obj_slack, together with scalar statistics alpha, lambda_w, N_act, N_fly, n_uavs, and "
    "n_tasks_active for the current BCD iteration."
)

LEVEL2_DECISION_RULES = [
    "The generated trajectory objective must assign dynamic_traj_objective using only positive scalar coefficients.",
    "Math transforms may be applied only to scalar values, not to CVXPY expressions.",
    "obj_slack must always be included with a positive coefficient.",
    "The output must not import cvxpy or gurobipy and must not define functions.",
]

LEVEL2_OUTPUT = (
    "Output only a JSON object with traj_obj_description and traj_obj_code. "
    "The traj_obj_code must directly assign dynamic_traj_objective."
)


PROMPT_BOXES = [
    {
        "title": "Level-1 Prompt Template (Way 1: Random Explore)",
        "task_goal": LEVEL1_COMMON_TASK_GOAL,
        "environment_state": LEVEL1_COMMON_ENV,
        "reference_examples": "None. This operator explores a completely new objective without historical best-individual context.",
        "decision_intro": "The generated objective must satisfy the following operator-specific constraints:",
        "decision_rules": LEVEL1_DECISION_RULES + [
            "Choose exactly one dominant design bias from: delay-first, energy-first, deadline-risk-aware, load-balancing-aware, or proximity-aware.",
            "Prefer novelty over safety, but keep the formulation solver-compatible and clearly biased toward one main direction.",
        ],
        "output_spec": LEVEL1_OUTPUT,
    },
    {
        "title": "Level-1 Prompt Template (Way 2: Local Refine)",
        "task_goal": LEVEL1_COMMON_TASK_GOAL,
        "environment_state": (
            LEVEL1_COMMON_ENV
            + " In addition, you are given a Previous Best Run block with evaluated system cost, "
            "solver feedback, and the best objective code from earlier iterations."
        ),
        "reference_examples": (
            "Previous Best Run:\n"
            "Evaluated System Cost: {score}\n"
            "Solver Feedback: {solver_feedback}\n"
            "Objective Code:\n"
            "{obj_code}"
        ),
        "decision_intro": "The refined objective must satisfy the following operator-specific constraints:",
        "decision_rules": LEVEL1_DECISION_RULES + [
            "Preserve the main structural idea of the previous best objective.",
            "Repair exactly one weakness, such as coefficient emphasis, normalization, one penalty term, or one local-vs-offload preference rule.",
            "Do not redesign the whole objective from scratch, and do not add more than one new penalty component.",
        ],
        "output_spec": LEVEL1_OUTPUT,
    },
    {
        "title": "Level-1 Prompt Template (Way 3: Structural Mutate)",
        "task_goal": LEVEL1_COMMON_TASK_GOAL,
        "environment_state": (
            LEVEL1_COMMON_ENV
            + " You are also given a Previous Best Run block as a reference, but the new candidate "
            "must intentionally explore a structurally different decomposition strategy."
        ),
        "reference_examples": (
            "Previous Best Run:\n"
            "Evaluated System Cost: {score}\n"
            "Solver Feedback: {solver_feedback}\n"
            "Objective Code:\n"
            "{obj_code}"
        ),
        "decision_intro": "The mutated objective must satisfy the following operator-specific constraints:",
        "decision_rules": LEVEL1_DECISION_RULES + [
            "Change at least two aspects among dominant bias, cost decomposition, normalization strategy, penalty structure, and local-vs-offload preference logic.",
            "The candidate should feel structurally different rather than cosmetically different.",
            "You may sacrifice some strengths of the previous best candidate if that helps explore a new direction.",
        ],
        "output_spec": LEVEL1_OUTPUT,
    },
    {
        "title": "Level-1 Prompt Template (Way 4: Resource-Aware Specialization)",
        "task_goal": LEVEL1_COMMON_TASK_GOAL,
        "environment_state": (
            LEVEL1_COMMON_ENV
            + " Resource-aware guidance is additionally provided about energy budget sensitivity, "
            "load balancing, deadline-energy tradeoff, and future-slot reservation."
        ),
        "reference_examples": (
            "Resource-Aware Guidance:\n"
            "1. Penalize assignments when cumulative UAV energy usage approaches E_max.\n"
            "2. Avoid concentrating all tasks on one UAV; encourage load balancing.\n"
            "3. Let urgent tasks prefer the nearest UAV even if energy cost is higher.\n"
            "4. Reserve future capacity when a UAV is heavily loaded in early slots."
        ),
        "decision_intro": "The specialized objective must satisfy the following operator-specific constraints:",
        "decision_rules": LEVEL1_DECISION_RULES + [
            "Choose exactly one dominant resource bias from: energy-budget-protection, load-balancing-pressure, urgent-task-nearest-UAV, or future-slot-reservation.",
            "Build one clearly identifiable term around that chosen resource bias rather than producing a generic balanced objective.",
        ],
        "output_spec": LEVEL1_OUTPUT,
    },
    {
        "title": "Level-2b Prompt Template (Way 1: New Weighting Strategy)",
        "task_goal": LEVEL2_COMMON_TASK_GOAL,
        "environment_state": LEVEL2_COMMON_ENV,
        "reference_examples": "No previous trajectory code is provided. The operator explores a completely new weighting strategy.",
        "decision_intro": "The generated weighting strategy must satisfy the following constraints:",
        "decision_rules": LEVEL2_DECISION_RULES + [
            "You may adapt weights using scalar transforms of N_act, N_fly, or their combinations.",
            "Think creatively about when communication should dominate and when propulsion should be relaxed.",
        ],
        "output_spec": LEVEL2_OUTPUT,
    },
    {
        "title": "Level-2b Prompt Template (Way 2: Improve Current Best)",
        "task_goal": LEVEL2_COMMON_TASK_GOAL,
        "environment_state": (
            LEVEL2_COMMON_ENV
            + " The current best trajectory objective code is also provided for targeted refinement."
        ),
        "reference_examples": "Current Best Trajectory Objective:\n{best_traj_code}",
        "decision_intro": "The improved weighting strategy must satisfy the following constraints:",
        "decision_rules": LEVEL2_DECISION_RULES + [
            "Keep at least 50% of the original structure.",
            "Refine scalar coefficients or the way N_act and N_fly affect the weights.",
            "A slightly different math transform is allowed, but the result should remain an improvement rather than a rewrite.",
        ],
        "output_spec": LEVEL2_OUTPUT,
    },
    {
        "title": "Level-2b Prompt Template (Way 3: Structural Change)",
        "task_goal": LEVEL2_COMMON_TASK_GOAL,
        "environment_state": (
            LEVEL2_COMMON_ENV
            + " The current best trajectory objective is given only as a reference for designing a structurally different strategy."
        ),
        "reference_examples": "Current Best Trajectory Objective (for reference only):\n{best_traj_code}",
        "decision_intro": "The structurally different weighting strategy must satisfy the following constraints:",
        "decision_rules": LEVEL2_DECISION_RULES + [
            "Change at least two aspects among math transform family, adaptive driver parameter, and communication-vs-propulsion emphasis.",
            "Do not copy the reference code; produce a genuinely distinct approach.",
        ],
        "output_spec": LEVEL2_OUTPUT,
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "9A9A9A", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = tc_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        child = tc_mar.find(qn(f"w:{name}"))
        if child is None:
            child = OxmlElement(f"w:{name}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def style_run(run, *, bold=False, italic=False, color=None, size=12) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def format_paragraph(paragraph, space_after=4) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.15


def add_labeled_paragraph(cell, label: str, body: str, italic_body: bool = False) -> None:
    p = cell.add_paragraph()
    format_paragraph(p)
    style_run(p.add_run(label), bold=True)
    style_run(p.add_run(body), italic=italic_body)


def add_multiline_reference(cell, title: str, body: str) -> None:
    p = cell.add_paragraph()
    format_paragraph(p)
    style_run(p.add_run(title), bold=True)
    for idx, line in enumerate(body.splitlines()):
        if idx == 0:
            p.add_run(" ")
        else:
            p.add_run().add_break(WD_BREAK.LINE)
        style_run(p.add_run(line), italic=("{Example_set}" in line or "{obj_code}" in line or "{best_traj_code}" in line))


def add_numbered_rule(cell, text: str) -> None:
    p = cell.add_paragraph(style="List Number")
    format_paragraph(p)
    style_run(p.add_run(text))


def add_prompt_box(doc: Document, spec: dict[str, object]) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.8)

    title_cell = table.cell(0, 0)
    body_cell = table.cell(1, 0)

    for cell in (title_cell, body_cell):
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_cell_shading(title_cell, "8C8C8C")
    set_cell_shading(body_cell, "ECECEC")

    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style_run(title_p.add_run(str(spec["title"])), bold=True, color=RGBColor(255, 255, 255), size=13)

    body_cell.paragraphs[0].clear()
    add_labeled_paragraph(body_cell, "Task goal: ", str(spec["task_goal"]))
    add_labeled_paragraph(body_cell, "Environment state: ", str(spec["environment_state"]))
    add_multiline_reference(body_cell, "Reference examples:", str(spec["reference_examples"]))
    add_labeled_paragraph(body_cell, "Decision rules: ", str(spec["decision_intro"]))
    for rule in spec["decision_rules"]:
        add_numbered_rule(body_cell, str(rule))
    add_labeled_paragraph(body_cell, "Output specification: ", str(spec["output_spec"]))


def build_docx() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    format_paragraph(heading, space_after=8)
    style_run(heading.add_run("Prompt Operator Templates"), bold=True, size=15)

    intro = doc.add_paragraph()
    format_paragraph(intro, space_after=10)
    style_run(
        intro.add_run(
            "This file summarizes all current prompt operators on the main Edge-UAV chain. "
            "Each operator is shown as an independent paper-style prompt box."
        ),
        size=11,
    )

    for idx, spec in enumerate(PROMPT_BOXES):
        add_prompt_box(doc, spec)
        if idx != len(PROMPT_BOXES) - 1:
            spacer = doc.add_paragraph()
            format_paragraph(spacer, space_after=4)

    return doc


def main() -> int:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_docx()
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
